import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import io
import zipfile
from datetime import datetime
import re

# --- 頁面設定 ---
st.set_page_config(page_title="銀行驗收單生成器-穩定增強版", page_icon="🏦", layout="wide")
st.title("🏦 銀行驗收單自動生成系統 (v2.7)")

# --- 函式：設定字體為標楷體 ---
def set_font_kai(run):
    try:
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    except:
        pass

# --- 函式：替換文字並保留字體格式 (強化防錯版) ---
def replace_text_in_document(doc, replacements):
    # 處理段落
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip(): continue # 跳過空段落
        for key, value in replacements.items():
            if key in paragraph.text:
                try:
                    full_text = paragraph.text.replace(key, str(value))
                    # 清空並重寫以維持標楷體
                    for run in paragraph.runs:
                        run.text = ""
                    new_run = paragraph.add_run(full_text)
                    set_font_kai(new_run)
                except Exception:
                    paragraph.text = paragraph.text.replace(key, str(value))
    
    # 處理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip(): continue
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            try:
                                new_text = paragraph.text.replace(key, str(value))
                                for run in paragraph.runs:
                                    run.text = ""
                                new_run = paragraph.add_run(new_text)
                                set_font_kai(new_run)
                            except Exception:
                                paragraph.text = paragraph.text.replace(key, str(value))

# --- 函式：手動日期解析 ---
def parse_date(date_str):
    date_str = str(date_str).strip()
    clean_date = re.sub(r'[^0-9]', '', date_str)
    if len(clean_date) == 8:
        try:
            return datetime.strptime(clean_date, "%Y%m%d").date()
        except:
            return None
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try:
            return datetime.strptime(date_str, fmt).date()
        except:
            continue
    return None

# --- 側邊欄 ---
st.sidebar.header("📁 檔案上傳")
uploaded_excel = st.sidebar.file_uploader("1. 上傳 Excel 清單 (.xlsx)", type=['xlsx'])
uploaded_word = st.sidebar.file_uploader("2. 上傳 Word 範本 (.docx)", type=['docx'])

if uploaded_excel and uploaded_word:
    # 讀取 Excel
    df = pd.read_excel(uploaded_excel, dtype=str)
    df.columns = [str(c).strip() for c in df.columns] # 清理所有標題空格
    
    # --- 智慧欄位尋找邏輯 ---
    cols = df.columns.tolist()
    # 鎖定 C 欄日期 (移除萊爾富)
    date_col = next((c for c in cols if "移除萊爾富" in c), None) or next((c for c in cols if "汰換日期" in c), cols[0])
    # 鎖定機號
    mid_col = next((c for c in cols if "機號" in c), "機號")
    # 鎖定財編
    aid_col = next((c for c in cols if "財編" in c or "Asset" in c), "CUB財編")
    
    st.sidebar.success(f"已辨識日期欄位：{date_col}")
    
    df['日期物件'] = pd.to_datetime(df[date_col], errors='coerce')
    
    st.header("⚙️ 篩選與產出設定")
    col1, col2 = st.columns(2)
    
    with col1:
        # 工程師名單處理
        eng_col = next((c for c in cols if "工程師" in c), "工程師")
        raw_list = df[eng_col].dropna().unique().tolist()
        all_engineers = sorted([str(e).strip() for e in raw_list if str(e).strip().lower() != 'nan' and str(e).strip() != ''])
        selected_engineers = st.multiselect("選擇需要的工程師：", options=all_engineers)
        
    with col2:
        date_mode = st.radio("日期選擇方式：", ["日曆選擇器", "手動輸入區間"], horizontal=True)
        start_date, end_date = None, None
        if date_mode == "日曆選擇器":
            valid_dates = df['日期物件'].dropna()
            if not valid_dates.empty:
                dr = st.date_input("選擇日期區區間：", [valid_dates.min().date(), valid_dates.max().date()])
                if len(dr) == 2: start_date, end_date = dr
        else:
            c1, c2 = st.columns(2)
            s_i = c1.text_input("開始(YYYYMMDD)", "")
            e_i = c2.text_input("結束(YYYYMMDD)", "")
            if s_i and e_i:
                start_date, end_date = parse_date(s_i), parse_date(e_i)

    # 過濾
    if start_date and end_date and selected_engineers:
        mask = (df[eng_col].astype(str).str.strip().isin(selected_engineers)) & \
               (df['日期物件'].dt.date >= start_date) & \
               (df['日期物件'].dt.date <= end_date)
        final_df = df[mask]
    else:
        final_df = pd.DataFrame()

    st.write(f"📊 目前篩選條件下共有 **{len(final_df)}** 筆資料。")

    if st.button("🚀 開始生成驗收單") and not final_df.empty:
        zip_buffer = io.BytesIO()
        progress_bar = st.progress(0)
        
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for i, (idx, row) in enumerate(final_df.iterrows()):
                uploaded_word.seek(0)
                doc = Document(uploaded_word)
                
                def gv(name, default_col):
                    # 優先從自動辨識的欄位抓取
                    actual_col = next((c for c in cols if name in c), default_col)
                    val = row.get(actual_col, '')
                    return "" if pd.isna(val) or str(val).lower() == 'nan' else str(val).strip()

                m_id = row.get(mid_col, '')
                if pd.isna(m_id) or str(m_id).lower() == 'nan' or not str(m_id).strip(): m_id = "(缺機號)"
                
                a_id = row.get(aid_col, '')
                if pd.isna(a_id) or str(a_id).lower() == 'nan' or not str(a_id).strip(): a_id = "(缺財編)"

                raw_dt = row.get('日期物件')
                fmt_dt = raw_dt.strftime("%Y年%m月%d日") if pd.notna(raw_dt) else "(日期缺失)"
                
                is_4g = gv('4G', '4G')
                if "無" in is_4g:
                    sim, ip, model = "", "", "FortiGate40F"
                else:
                    sim, ip, model = gv('SIM卡編號', 'SIM卡編號'), gv('SIM卡IP', 'SIM卡IP'), "FortiGate40F 3G/4G"
                
                replacements = {
                    "{{Date}}": fmt_dt,
                    "{{Station}}": gv('站點名稱', '站點名稱'),
                    "{{MachineID}}": m_id,
                    "{{Address}}": gv('地址', '地址'),
                    "{{SN}}": gv('機器序號', '機器序號'),
                    "{{AssetID}}": a_id,
                    "{{SIM}}": sim,
                    "{{IP}}": ip,
                    "{{Model}}": model
                }
                
                replace_text_in_document(doc, replacements)
                
                doc_io = io.BytesIO()
                doc.save(doc_io)
                zip_file.writestr(f"{'[Error]' if '(缺' in m_id or '(缺' in a_id else ''}{raw_dt.strftime('%Y%m%d') if pd.notna(raw_dt) else 'NoDate'}_{m_id}_{gv('站點名稱', '站點名稱').replace('/', '_')}.docx", doc_io.getvalue())
                progress_bar.progress((i + 1) / len(final_df))
        
        st.success("✅ 生成完成！")
        st.download_button("📥 下載 ZIP 檔案", zip_buffer.getvalue(), f"驗收單_{datetime.now().strftime('%Y%m%d_%H%M')}.zip")
else:
    st.info("請上傳檔案。")
